# -*- coding: utf-8 -*-
"""
DocumentParser — Extracts text/data from various document types.

Supported:
- PDF  → PyMuPDF (fitz) text extraction + Gemini Vision for scanned pages
- DOCX → python-docx
- XLSX → openpyxl
- Images (PNG, JPG, JPEG) → Gemini Vision OCR
"""
import logging
import io

_logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parses various document types into extractable text.
    For images and scanned PDFs, delegates to GeminiService vision OCR.
    """

    IMAGE_MIMES = {'image/png', 'image/jpeg', 'image/jpg', 'image/webp', 'image/gif'}
    PDF_MIME = 'application/pdf'
    DOCX_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    XLSX_MIME = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    XLS_MIME = 'application/vnd.ms-excel'

    def __init__(self, gemini_service=None):
        """
        :param gemini_service: GeminiService instance for vision OCR (optional but recommended)
        """
        self.gemini_service = gemini_service

    def parse(self, file_bytes, filename='', mimetype=''):
        """
        Auto-detect file type and extract text.

        :param file_bytes: raw bytes of the file
        :param filename: original filename (used for type detection fallback)
        :param mimetype: MIME type string
        :return: str — extracted text content
        """
        mimetype_lower = (mimetype or '').lower()
        filename_lower = (filename or '').lower()

        if mimetype_lower in self.IMAGE_MIMES or any(
            filename_lower.endswith(ext) for ext in ('.png', '.jpg', '.jpeg', '.webp')
        ):
            text = self.parse_image(file_bytes, mimetype_lower or 'image/jpeg')
        elif mimetype_lower == self.PDF_MIME or filename_lower.endswith('.pdf'):
            text = self.parse_pdf(file_bytes)
        elif mimetype_lower == self.DOCX_MIME or filename_lower.endswith('.docx'):
            text = self.parse_docx(file_bytes)
        elif mimetype_lower == self.XLS_MIME or filename_lower.endswith('.xls'):
            text = self.parse_xls(file_bytes)
        elif mimetype_lower == self.XLSX_MIME or filename_lower.endswith('.xlsx'):
            text = self.parse_xlsx(file_bytes)
        else:
            # Fallback: treat as plain text
            try:
                text = file_bytes.decode('utf-8', errors='replace')
            except Exception:
                text = ''

        return self._clean_and_filter_text(text)

    def _clean_and_filter_text(self, text):
        """
        Clean up extracted text by:
        1. Compressing multiple blank lines into a maximum of 2 newlines.
        2. Compressing large gaps of spaces (3 or more) into exactly 2 spaces to preserve tabular alignment while saving tokens.
        3. Stripping trailing spaces from each line while preserving leading indentation.
        """
        if not text:
            return ''
        import re
        
        # Strip trailing whitespaces from each line and compress space gaps
        lines = []
        for line in text.split('\n'):
            line_rstripped = line.rstrip()
            if line_rstripped.strip():
                # Compress 3 or more consecutive spaces preceded by a non-space char into exactly 2 spaces
                line_cleaned = re.sub(r'(?<=\S) {3,}', '  ', line_rstripped)
                lines.append(line_cleaned)
            else:
                lines.append('')
        
        cleaned_text = '\n'.join(lines)
        
        # Compress multiple blank lines (3 or more consecutive newlines) into a maximum of 2 newlines
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        return cleaned_text.strip('\n')

    def _get_fitz(self):
        """Import PyMuPDF — supports both old (fitz) and new (pymupdf) import names."""
        try:
            import fitz
            return fitz
        except ImportError:
            pass
        try:
            import pymupdf as fitz
            return fitz
        except ImportError:
            pass
        return None

    def parse_pdf(self, file_bytes):
        """
        Extract text from PDF using PyMuPDF (fitz or pymupdf import).
        Falls back to PyPDF2, then to Groq Vision OCR for scanned PDFs.
        """
        fitz = self._get_fitz()
        if fitz is None:
            _logger.warning('PyMuPDF not found. Trying PyPDF2 fallback.')
            return self._parse_pdf_pypdf2(file_bytes)

        texts = []
        try:
            doc = fitz.open(stream=file_bytes, filetype='pdf')
            for page_num, page in enumerate(doc):
                page_text = page.get_text('text').strip()
                if page_text and len(page_text) > 50:
                    texts.append(f'--- Page {page_num + 1} ---\n{page_text}')
                elif self.gemini_service:
                    # Scanned page — render and send to vision
                    _logger.info('Page %d appears scanned, using vision OCR', page_num + 1)
                    pix = page.get_pixmap(dpi=200)
                    img_bytes = pix.tobytes('jpeg')
                    ocr_text = self.gemini_service.vision_ocr(img_bytes, 'image/jpeg')
                    texts.append(f'--- Page {page_num + 1} (OCR) ---\n{ocr_text}')
            doc.close()
        except Exception as e:
            _logger.error('PDF parsing error: %s', str(e))
            return self._vision_fallback(file_bytes, 'application/pdf')

        return '\n\n'.join(texts)

    def _parse_pdf_pypdf2(self, file_bytes):
        """Fallback PDF text extraction using PyPDF2."""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            texts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ''
                if text.strip():
                    texts.append(f'--- Page {i + 1} ---\n{text.strip()}')
                elif self.gemini_service:
                    _logger.info('PyPDF2: page %d is empty, using vision OCR', i + 1)
            return '\n\n'.join(texts) if texts else self._vision_fallback(file_bytes, 'application/pdf')
        except ImportError:
            _logger.warning('PyPDF2 not installed either. Falling back to vision OCR.')
            return self._vision_fallback(file_bytes, 'application/pdf')
        except Exception as e:
            _logger.error('PyPDF2 parsing error: %s', str(e))
            return self._vision_fallback(file_bytes, 'application/pdf')

    def parse_docx(self, file_bytes):
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            _logger.warning('python-docx not installed. Install with: pip install python-docx')
            return '[DOCX parsing unavailable - install python-docx]'

        try:
            doc = Document(io.BytesIO(file_bytes))
            parts = []

            # Paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(' | '.join(cells))
                if table_rows:
                    parts.append('\n'.join(table_rows))

            return '\n'.join(parts)
        except Exception as e:
            _logger.error('DOCX parsing error: %s', str(e))
            return f'[Error parsing DOCX: {str(e)}]'

    def parse_xlsx(self, file_bytes):
        """Extract text from Excel using openpyxl."""
        try:
            import openpyxl
        except ImportError:
            _logger.warning('openpyxl not installed. Install with: pip install openpyxl')
            return '[Excel parsing unavailable - install openpyxl]'

        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f'=== Sheet: {sheet_name} ===')
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(cell) if cell is not None else '' for cell in row]
                    row_text = ' | '.join(row_vals).strip(' |')
                    if row_text.strip():
                        parts.append(row_text)
            return '\n'.join(parts)
        except Exception as e:
            _logger.error('Excel parsing error: %s', str(e))
            return f'[Error parsing Excel: {str(e)}]'

    def parse_xls(self, file_bytes):
        """Extract text from old Excel (.xls) using xlrd."""
        try:
            import xlrd
        except ImportError:
            _logger.warning('xlrd not installed. Install with: pip install xlrd')
            return '[Excel .xls parsing unavailable - install xlrd]'

        try:
            book = xlrd.open_workbook(file_contents=file_bytes)
            parts = []
            for sheet_index in range(book.nsheets):
                sheet = book.sheet_by_index(sheet_index)
                parts.append(f'=== Sheet: {sheet.name} ===')
                for row_idx in range(sheet.nrows):
                    row_vals = []
                    for col_idx in range(sheet.ncols):
                        cell = sheet.cell_value(row_idx, col_idx)
                        row_vals.append(str(cell) if cell is not None and cell != '' else '')
                    row_text = ' | '.join(row_vals).strip(' |')
                    if row_text.strip():
                        parts.append(row_text)
            return '\n'.join(parts)
        except Exception as e:
            _logger.error('Excel .xls parsing error: %s', str(e))
            return f'[Error parsing .xls Excel: {str(e)}]'

    def parse_image(self, image_bytes, mime_type='image/jpeg'):
        """Use Gemini Vision to extract text from an image."""
        if not self.gemini_service:
            return '[Image OCR unavailable - Gemini service not configured]'
        try:
            return self.gemini_service.vision_ocr(image_bytes, mime_type)
        except Exception as e:
            _logger.error('Image OCR error: %s', str(e))
            return f'[Image OCR error: {str(e)}]'

    def _vision_fallback(self, file_bytes, mime_type):
        """Generic vision fallback when text parsing fails."""
        if not self.gemini_service:
            return '[Document parsing failed and vision OCR unavailable]'
        try:
            # For non-image files, try rendering first page as image
            return self.gemini_service.vision_ocr(file_bytes[:1024 * 1024], 'image/jpeg')
        except Exception as e:
            return f'[Document parsing failed: {str(e)}]'

    def parse_attachment(self, attachment):
        """
        Parse an Odoo ir.attachment record.

        :param attachment: ir.attachment recordset (single)
        :return: str — extracted text
        """
        import base64
        try:
            if attachment.datas:
                file_bytes = base64.b64decode(attachment.datas)
            else:
                return ''
            return self.parse(
                file_bytes=file_bytes,
                filename=attachment.name or '',
                mimetype=attachment.mimetype or '',
            )
        except Exception as e:
            _logger.error('Attachment parsing error for %s: %s', attachment.name, str(e))
            return f'[Error reading file: {str(e)}]'
